"""
Payroll PDF Generator Views
"""

# Standard library imports
import io
import math
import os
import subprocess
import zipfile
from datetime import datetime, timedelta
from typing import Dict, Tuple

# Third-party imports
import stripe
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from django.conf import settings
from django.http import HttpRequest, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST

# Local application imports
from .models import PaymentToken

# Stripe configuration
stripe.api_key = settings.STRIPE_SECRET_KEY



# Constants
TAX_BRACKETS = [
    (0, 10275, 0.10),
    (10276, 41775, 0.12),
    (41776, 89075, 0.22),
    (89076, 170050, 0.24),
    (170051, 215950, 0.32),
    (215951, 539900, 0.35),
    (539901, float('inf'), 0.37)
]

SOCIAL_SECURITY_RATE = 0.062
MEDICARE_RATE = 0.0145

START_DATE = datetime(2023, 12, 21)
PAY_DATE_REFERENCE = datetime(2023, 12, 22)
BIWEEKLY_DAYS = 14
WEEKLY_DAYS = 7

NUMBER_WORDS = {
    0: "ZERO", 1: "ONE", 2: "TWO", 3: "THREE", 4: "FOUR", 5: "FIVE",
    6: "SIX", 7: "SEVEN", 8: "EIGHT", 9: "NINE", 10: "TEN",
    11: "ELEVEN", 12: "TWELVE", 13: "THIRTEEN", 14: "FOURTEEN", 15: "FIFTEEN",
    16: "SIXTEEN", 17: "SEVENTEEN", 18: "EIGHTEEN", 19: "NINETEEN",
    20: "TWENTY", 30: "THIRTY", 40: "FORTY", 50: "FIFTY",
    60: "SIXTY", 70: "SEVENTY", 80: "EIGHTY", 90: "NINETY"
}

HUNDREDS_WORDS = {
    100: "HUNDRED",
    1000: "THOUSAND"
}


class PayrollCalculator:
    """Calculate payroll deductions and taxes"""
    
    @staticmethod
    def calculate(annual_salary: int, pay_period: int) -> Tuple[float, float, float, float, float]:
        """
        Calculate payroll values
        
        Args:
            annual_salary: Annual salary amount
            pay_period: Number of pay periods per year
            
        Returns:
            Tuple of (gross_salary, fed_withholding, ss, medicare, fica_deduction)
        """
        gross_salary = PayrollCalculator._round_up(annual_salary / pay_period)
        fed_withholding = PayrollCalculator._round_up(
            gross_salary * PayrollCalculator._get_tax_rate(annual_salary)
        )
        ss = PayrollCalculator._round_up(gross_salary * SOCIAL_SECURITY_RATE)
        medicare = PayrollCalculator._round_up(gross_salary * MEDICARE_RATE)
        fica_deduction = PayrollCalculator._round_up(fed_withholding + ss + medicare)
        
        return gross_salary, fed_withholding, ss, medicare, fica_deduction
    
    @staticmethod
    def _get_tax_rate(income: float) -> float:
        """Get applicable tax rate based on income"""
        for lower, upper, rate in TAX_BRACKETS:
            if lower <= income <= upper:
                return rate
        return TAX_BRACKETS[-1][2]  # Return highest bracket if not found
    
    @staticmethod
    def _round_up(number: float, decimals: int = 2) -> float:
        """Round up to specified decimal places"""
        factor = 10 ** decimals
        return math.ceil(number * factor) / factor


class NumberFormatter:
    """Format numbers for display"""
    
    @staticmethod
    def format_currency(number: float) -> str:
        """Format number as currency with commas and 2 decimals"""
        return "{:,.2f}".format(number)
    
    @staticmethod
    def get_decimal_part(number: float) -> str:
        """Extract decimal part of a number as string"""
        number_str = str(number)
        parts = number_str.split('.')
        if len(parts) > 1:
            return parts[1][:2].ljust(2, '0')
        return "00"
    
    @staticmethod
    def number_to_words(num: int) -> str:
        """Convert number to words in English (0-9999)"""
        if num < 20:
            return NUMBER_WORDS[num]
        elif num < 100:
            tens = num // 10 * 10
            ones = num % 10
            if ones == 0:
                return NUMBER_WORDS[tens]
            return f"{NUMBER_WORDS[tens]} {NUMBER_WORDS[ones]}"
        elif num < 1000:
            hundreds = num // 100
            remainder = num % 100
            if remainder == 0:
                return f"{NUMBER_WORDS[hundreds]} {HUNDREDS_WORDS[100]}"
            return f"{NUMBER_WORDS[hundreds]} {HUNDREDS_WORDS[100]} AND {NumberFormatter.number_to_words(remainder)}"
        elif num < 10000:
            thousands = num // 1000
            remainder = num % 1000
            if remainder == 0:
                return f"{NUMBER_WORDS[thousands]} {HUNDREDS_WORDS[1000]}"
            return f"{NUMBER_WORDS[thousands]} {HUNDREDS_WORDS[1000]} {NumberFormatter.number_to_words(remainder)}"
        return "Number out of range (must be between 0 and 9999)"


class PayDateCalculator:
    """Calculate and validate pay dates"""
    
    @staticmethod
    def get_correct_pay_date(pay_date: datetime) -> datetime:
        """
        Adjust pay date to fall on a valid biweekly schedule
        
        Args:
            pay_date: The requested pay date
            
        Returns:
            Adjusted pay date that falls on the biweekly schedule
        """
        if (pay_date - datetime(2024, 1, 5)).days % BIWEEKLY_DAYS == 0:
            return pay_date
        
        rounded = (pay_date - PAY_DATE_REFERENCE).days // BIWEEKLY_DAYS
        mult = rounded * BIWEEKLY_DAYS
        return PAY_DATE_REFERENCE + timedelta(days=mult)


class PayrollDocumentGenerator:
    """Generate payroll documents"""
    
    def __init__(self, request_data: Dict, payroll_data: Tuple):
        self.request_data = request_data
        self.gross_salary, self.fed_withholding, self.ss, self.medicare, self.fica_deduction = payroll_data
        self.output_dir = os.path.join('media', 'pdfs')
        
    def generate_multiple_pdfs(self) -> HttpResponse:
        """Generate multiple payroll PDFs and return as ZIP"""
        start_period = PayDateCalculator.get_correct_pay_date(
            datetime.strptime(self.request_data['start_period'], '%Y-%m-%d')
        )
        end_period = PayDateCalculator.get_correct_pay_date(
            datetime.strptime(self.request_data['end_period'], '%Y-%m-%d')
        )
        number_payments = (end_period - start_period).days // BIWEEKLY_DAYS
        period = int(self.request_data['period'])
        
        temp_files = []
        final_pdf_paths = []
        
        try:
            self._ensure_output_dir()
            
            for i in range(number_payments):
                start_period += timedelta(days=BIWEEKLY_DAYS)
                payment_number = self._calculate_payment_number(start_period, period)
                
                temp_docx, final_pdf = self._generate_single_pdf(
                    i, start_period, payment_number
                )
                temp_files.append(temp_docx)
                final_pdf_paths.append(final_pdf)
            
            return self._create_zip_response(final_pdf_paths)
            
        finally:
            self._cleanup_files(temp_files + final_pdf_paths)
    
    def _ensure_output_dir(self):
        """Create output directory if it doesn't exist"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def _calculate_payment_number(self, start_period: datetime, period: int) -> int:
        """Calculate the payment number based on period"""
        days_divisor = BIWEEKLY_DAYS if period == 26 else WEEKLY_DAYS
        return (start_period - START_DATE).days // days_divisor
    
    def _generate_single_pdf(self, index: int, start_period: datetime, 
                            payment_number: int) -> Tuple[str, str]:
        """Generate a single payroll PDF"""
        temp_docx_path = f'temp_modified_{index}.docx'
        
        # Create and modify document
        doc = Document('base.docx')
        replacements = self._get_replacements(start_period, payment_number)
        self._apply_replacements(doc, replacements)
        
        # Save temporary docx
        doc.save(temp_docx_path)
        
        # Convert to PDF
        final_pdf_path = self._convert_to_pdf(temp_docx_path, index, start_period)
        
        return temp_docx_path, final_pdf_path
    
    def _get_replacements(self, start_period: datetime, payment_number: int) -> Dict[str, str]:
        """Get dictionary of placeholder replacements"""
        net_pay = self.gross_salary - self.fica_deduction
        check_id = self.request_data.get('check_id', '')
        
        return {
            '<<nombre>>': f"{self.request_data['name']} {self.request_data['last_name']}",
            '<<client_address>>': self.request_data['client_address'],
            '<<company>>': self.request_data['company'],
            '<<city_state>>': self.request_data['city_state'],
            '<<address_co>>': self.request_data['address_co'],
            '<<check_id>>': str(check_id),
            '<<fecha>>': start_period.strftime('%m/%d/%Y'),
            '<<pay_date>>': (start_period - timedelta(days=BIWEEKLY_DAYS)).strftime('%m/%d/%Y'),
            '<<netpaytext>>': NumberFormatter.number_to_words(round(net_pay)),
            '<<decimal>>': NumberFormatter.get_decimal_part(net_pay),
            '<<ssn_digits>>': self.request_data['ssn_digits'],
            '<<netpay>>': NumberFormatter.format_currency(PayrollCalculator._round_up(net_pay)),
            '<<dependents>>': self.request_data['dependents'],
            '<<salary>>': NumberFormatter.format_currency(self.gross_salary),
            '<<fed>>': NumberFormatter.format_currency(self.fed_withholding),
            '<<ss>>': NumberFormatter.format_currency(self.ss),
            '<<mc>>': NumberFormatter.format_currency(self.medicare),
            '<<totalt>>': NumberFormatter.format_currency(
                PayrollCalculator._round_up(self.fed_withholding + self.ss + self.medicare)
            ),
            # Year to Date
            '<<salaryytd>>': NumberFormatter.format_currency(
                PayrollCalculator._round_up(self.gross_salary * payment_number)
            ),
            '<<fedytd>>': NumberFormatter.format_currency(
                PayrollCalculator._round_up(self.fed_withholding * payment_number)
            ),
            '<<ssytd>>': NumberFormatter.format_currency(
                PayrollCalculator._round_up(self.ss * payment_number)
            ),
            '<<mcytd>>': NumberFormatter.format_currency(
                PayrollCalculator._round_up(self.medicare * payment_number)
            ),
            '<<totaltytd>>': NumberFormatter.format_currency(
                PayrollCalculator._round_up(
                    (self.fed_withholding + self.ss + self.medicare) * payment_number
                )
            ),
        }
    
    def _apply_replacements(self, doc: Document, replacements: Dict[str, str]):
        """Apply text replacements and formatting to document"""
        no_font_size_changes = {'<<nombre>>', '<<fecha>>', '<<netpay>>', 
                               '<<dependents>>', '<<pay_date>>'}
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
        
        # Replace and format in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                            self._format_cell(cell, key, no_font_size_changes)
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    def _format_cell(self, cell, key: str, no_font_size_changes: set):
        """Format a cell based on the key"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if key not in no_font_size_changes:
                    run.font.size = Pt(7)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif key == '<<netpay>>':
                    run.font.size = Pt(9)
                    run.bold = True
                else:
                    run.font.size = Pt(8)
    
    def _convert_to_pdf(self, docx_path: str, index: int, 
                       start_period: datetime) -> str:
        """Convert DOCX to PDF using LibreOffice"""
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', 
             '--outdir', self.output_dir, docx_path],
            capture_output=True,
            text=True
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr}")
        
        temp_pdf_path = os.path.join(self.output_dir, f'temp_modified_{index}.pdf')
        if not os.path.exists(temp_pdf_path):
            raise FileNotFoundError(f"PDF not created: {temp_pdf_path}")
        
        # Rename to final name
        pdf_name = (f"{self.request_data['name']}{self.request_data['last_name']}_"
                   f"{start_period.strftime('%m%d%Y')}.pdf")
        final_pdf_path = os.path.join(self.output_dir, pdf_name)
        os.rename(temp_pdf_path, final_pdf_path)
        
        return final_pdf_path
    
    def _create_zip_response(self, pdf_paths: list) -> HttpResponse:
        """Create ZIP file response with all PDFs"""
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            for pdf_path in pdf_paths:
                zip_file.write(pdf_path, os.path.basename(pdf_path))
        
        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer, content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename="payroll_pdfs.zip"'
        return response
    
    @staticmethod
    def _cleanup_files(file_paths: list):
        """Remove temporary files"""
        for path in file_paths:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass  # Ignore cleanup errors


# Views
PRODUCT_ID = settings.PRODUCT_ID

def index(request):
    """Main page with payment button - displays the checkout URL."""
    if request.method == 'GET':
        # Create Stripe checkout session
        checkout_session = create_stripe_checkout_session(PRODUCT_ID)
        
        # Create token linked to Stripe session
        payment_token = PaymentToken.objects.create(
            stripe_session_id=checkout_session.id
        )

        return render(request, 'index.html', {
            'checkout_url': checkout_session.url,
            'session_id': checkout_session.id
        })

    elif request.method == 'POST':
        client_email = request.POST.get('email')
        session_id = request.POST.get('session_id')

        if not session_id:
            return redirect('index')

        try:
            payment_token = PaymentToken.objects.get(stripe_session_id=session_id)
            payment_token.customer_email = client_email
            payment_token.save()

            checkout_session = stripe.checkout.Session.retrieve(session_id)
            return redirect(checkout_session.url)

        except PaymentToken.DoesNotExist:
            return redirect('index')
        except Exception as e:
            return HttpResponse(f"Error: {str(e)}", status=500)

    return HttpResponse("Invalid request method", status=405)


def payment_success(request):
    """View displayed when payment is successful."""
    if request.method == 'GET':
        return render(request, 'payments/success.html')

    elif request.method == 'POST':
        email = request.POST.get('email')

        token = PaymentToken.objects.filter(
            customer_email=email,
        ).first()

        if not token or not token.is_valid():
            return redirect('index')

        return redirect('payroll', token=token.token)


def payment_cancel(request):
    """View displayed when user cancels the payment."""
    return render(request, 'payments/cancel.html')


@csrf_exempt
def payroll_view(request: HttpRequest, token: str) -> HttpResponse:

    token_obj = get_object_or_404(PaymentToken, token=token)

    if not token_obj or not token_obj.is_valid():
        return redirect('index')
    """Main view for payroll generation."""
    if request.method == 'POST':
        try:
            annual_salary = int(request.POST['anual'])
            pay_period = int(request.POST['period'])

            # Calculate payroll
            payroll_data = PayrollCalculator.calculate(annual_salary, pay_period)

            # Generate PDFs
            generator = PayrollDocumentGenerator(request.POST, payroll_data)

            token = PaymentToken.objects.get(token=token)
            token.mark_as_used()

            return generator.generate_multiple_pdfs()

        except (ValueError, KeyError) as e:
            return HttpResponse(f"Invalid input data: {str(e)}", status=400)
        except RuntimeError as e:
            return HttpResponse(f"Error generating PDFs: {str(e)}", status=500)
    else:
        return render(request, 'payroll_view.html')


@csrf_exempt
@require_POST
def stripe_webhook(request):
    """Webhook for receiving Stripe events."""
    payload = request.body
    sig_header = request.META.get('HTTP_STRIPE_SIGNATURE')

    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, settings.STRIPE_WEBHOOK_SECRET
        )
    except ValueError:
        return HttpResponse(status=400)
    except stripe.error.SignatureVerificationError:
        return HttpResponse(status=400)

    # Handle successful checkout event
    if event['type'] == 'checkout.session.completed':
        session = event['data']['object']

        try:
            payment_token = PaymentToken.objects.get(
                stripe_session_id=session['id']
            )

            if session['payment_status'] == 'paid':
                payment_token.is_paid = True
                payment_token.paid_at = timezone.now()
                payment_token.stripe_payment_intent = session.get('payment_intent')

                if session.get('customer_details', {}).get('email'):
                    payment_token.customer_email = session['customer_details']['email']

                payment_token.save()
        except PaymentToken.DoesNotExist:
            pass

    return HttpResponse(status=200)


# Utility functions

def create_stripe_checkout_session(price_id):
    """Create a Stripe checkout session."""
    try:
        checkout_session = stripe.checkout.Session.create(
            line_items=[
                {
                    'price': price_id,
                    'quantity': 1,
                },
            ],
            mode='payment',
            success_url=f"{settings.DOMAIN}/payment/success/",
            cancel_url=f"{settings.DOMAIN}/payment/error/",
            automatic_tax={'enabled': True}
        )
        return checkout_session
    except Exception as e:
        raise Exception(f"Unexpected error: {str(e)}")
